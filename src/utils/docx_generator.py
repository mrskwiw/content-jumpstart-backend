"""DOCX deliverable generator for professional client documents

Uses lazy imports for docx library to reduce startup time.
Heavy modules are only loaded when DOCX generation is actually needed.
"""

from datetime import datetime
from pathlib import Path
from typing import TYPE_CHECKING, List, Optional

# Lazy imports - these modules are heavy and only loaded when needed
if TYPE_CHECKING:
    from docx import Document

from ..models.client_brief import ClientBrief
from ..models.post import Post
from ..models.qa_report import QAReport
from .logger import logger


class DOCXGenerator:
    """Generates professional DOCX deliverables for clients

    Uses lazy imports to load python-docx only when generation is needed.
    This reduces memory footprint and startup time when DOCX export is not used.
    """

    def __init__(self):
        """Initialize DOCX generator with default styles"""
        # Import RGBColor only when needed
        from docx.shared import RGBColor

        self.brand_color = RGBColor(41, 128, 185)  # Professional blue
        self.metadata_color = RGBColor(127, 140, 141)  # Gray for metadata

    def create_deliverable_docx(
        self,
        posts: List[Post],
        client_brief: ClientBrief,
        output_path: Path,
        include_voice_guide: bool = True,
        include_schedule: bool = True,
        qa_report: Optional[QAReport] = None,
        voice_guide_content: Optional[str] = None,
        schedule_content: Optional[str] = None,
    ) -> Path:
        """
        Generate complete DOCX deliverable

        Document Structure:
        - Cover page (company name, date, logo placeholder)
        - Introduction section (brief overview)
        - 30 posts (numbered, with template metadata)
        - Brand voice guide appendix (if included)
        - Posting schedule appendix (if included)
        - QA summary appendix (if provided)

        Args:
            posts: List of generated posts
            client_brief: Client brief information
            output_path: Path for output DOCX file
            include_voice_guide: Include voice guide section
            include_schedule: Include posting schedule section
            qa_report: Optional QA report to include
            voice_guide_content: Optional pre-generated voice guide markdown
            schedule_content: Optional pre-generated schedule markdown

        Returns:
            Path to generated DOCX file
        """
        logger.info(f"Generating DOCX deliverable for {client_brief.company_name}")

        # Lazy import heavy docx modules
        from docx import Document

        # Create document
        doc = Document()

        # Set up document styles
        self._configure_document_styles(doc)

        # Add sections
        self._add_cover_page(doc, client_brief)
        self._add_introduction(doc, client_brief)
        self._add_posts_section(doc, posts, client_brief)

        if include_voice_guide and voice_guide_content:
            self._add_voice_guide_section(doc, voice_guide_content)

        if include_schedule and schedule_content:
            self._add_schedule_section(doc, schedule_content)

        if qa_report:
            self._add_qa_summary_section(doc, qa_report)

        # Save document
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))

        logger.info(f"DOCX deliverable saved to {output_path}")
        return output_path

    def _configure_document_styles(self, doc: "Document"):
        """Configure custom styles for the document"""
        # Lazy import docx types
        from docx.enum.style import WD_STYLE_TYPE
        from docx.shared import Pt

        styles = doc.styles

        # Heading 1: Company/Section Headers
        if "Custom Heading 1" not in styles:
            h1_style = styles.add_style("Custom Heading 1", WD_STYLE_TYPE.PARAGRAPH)
            h1_style.font.name = "Arial"
            h1_style.font.size = Pt(16)
            h1_style.font.bold = True
            h1_style.font.color.rgb = self.brand_color

        # Heading 2: Post Titles
        if "Custom Heading 2" not in styles:
            h2_style = styles.add_style("Custom Heading 2", WD_STYLE_TYPE.PARAGRAPH)
            h2_style.font.name = "Arial"
            h2_style.font.size = Pt(14)
            h2_style.font.bold = True
            h2_style.font.color.rgb = self.brand_color

        # Metadata: Post metadata
        if "Metadata" not in styles:
            meta_style = styles.add_style("Metadata", WD_STYLE_TYPE.PARAGRAPH)
            meta_style.font.name = "Calibri"
            meta_style.font.size = Pt(9)
            meta_style.font.italic = True
            meta_style.font.color.rgb = self.metadata_color

    def _add_cover_page(self, doc: "Document", client_brief: ClientBrief):
        """Add professional cover page"""
        # Lazy import docx types
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt

        # Title
        title = doc.add_paragraph()
        title_run = title.add_run("30-Day Content Jumpstart")
        title_run.font.name = "Arial"
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title_run.font.color.rgb = self.brand_color
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Spacing
        doc.add_paragraph()

        # Client name
        client = doc.add_paragraph()
        client_run = client.add_run(client_brief.company_name)
        client_run.font.name = "Arial"
        client_run.font.size = Pt(22)
        client.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Spacing
        doc.add_paragraph()
        doc.add_paragraph()

        # Date
        date_para = doc.add_paragraph()
        date_run = date_para.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
        date_run.font.name = "Calibri"
        date_run.font.size = Pt(12)
        date_run.font.italic = True
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Page break
        doc.add_page_break()

    def _add_introduction(self, doc: "Document", client_brief: ClientBrief):
        """Add introduction section with client context"""
        # Section header
        intro_header = doc.add_heading("About This Content Package", level=1)

        # Client overview
        doc.add_paragraph(
            f"This content package has been custom-generated for {client_brief.company_name}. "
            f"All 30 posts are tailored to your brand voice, target audience, and business goals."
        )

        doc.add_paragraph()

        # Client details table
        table = doc.add_table(rows=5, cols=2)
        table.style = "Light Grid Accent 1"

        # Company
        table.rows[0].cells[0].text = "Company"
        table.rows[0].cells[1].text = client_brief.company_name

        # Business description
        table.rows[1].cells[0].text = "Business"
        table.rows[1].cells[1].text = client_brief.business_description

        # Target audience
        table.rows[2].cells[0].text = "Target Audience"
        table.rows[2].cells[1].text = client_brief.ideal_customer

        # Main problem solved
        table.rows[3].cells[0].text = "Main Problem Solved"
        table.rows[3].cells[1].text = client_brief.main_problem_solved

        # Target platforms
        table.rows[4].cells[0].text = "Target Platforms"
        platforms_text = ", ".join([p.value for p in client_brief.target_platforms])
        table.rows[4].cells[1].text = platforms_text

        doc.add_paragraph()

        # Usage instructions
        doc.add_heading("How to Use These Posts", level=2)

        usage_points = [
            "Review each post and customize as needed before publishing",
            "Use the posting schedule (see appendix) to plan your content calendar",
            "Track performance using the analytics tracker spreadsheet",
            "Refer to the brand voice guide when creating future content",
            "Feel free to combine or split posts based on your needs",
        ]

        for point in usage_points:
            para = doc.add_paragraph(point, style="List Bullet")

        # Page break
        doc.add_page_break()

    def _add_posts_section(self, doc: "Document", posts: List[Post], client_brief: ClientBrief):
        """Add all posts with formatting"""
        # Section header
        posts_header = doc.add_heading("Your 30 Posts", level=1)

        doc.add_paragraph(
            f"Below are your 30 custom posts for {', '.join([p.value for p in client_brief.target_platforms])}. "
            "Each post includes the template name and key metadata."
        )

        doc.add_paragraph()

        # Add each post
        for i, post in enumerate(posts, 1):
            self._add_post_entry(doc, post, i)

            # Add spacing between posts (but not after last post)
            if i < len(posts):
                doc.add_paragraph()

            # Page break every 5 posts for readability
            if i % 5 == 0 and i < len(posts):
                doc.add_page_break()

    def _add_post_entry(self, doc: "Document", post: Post, post_number: int):
        """Add single post with formatting"""
        # Lazy import docx types
        from docx.shared import Pt

        # Post header
        post_header = doc.add_heading(f"Post {post_number}: {post.template_name}", level=2)

        # Post content
        content_para = doc.add_paragraph(post.content)
        content_para.paragraph_format.space_after = Pt(12)

        # Metadata box
        metadata_parts = [
            f"Template: {post.template_name}",
            f"Words: {post.word_count}",
            f"Has CTA: {'Yes' if post.has_cta else 'No'}",
        ]

        # Add platform if specified
        if post.target_platform:
            metadata_parts.insert(1, f"Platform: {post.target_platform}")

        # Add keywords if present
        keywords = getattr(post, "keywords_used", [])
        if keywords:
            keywords_str = ", ".join(keywords[:3])  # First 3 keywords
            if len(keywords) > 3:
                keywords_str += f" (+{len(keywords) - 3} more)"
            metadata_parts.append(f"Keywords: {keywords_str}")

        metadata_text = " | ".join(metadata_parts)
        metadata_para = doc.add_paragraph(metadata_text, style="Metadata")

        # Horizontal line separator
        doc.add_paragraph("_" * 80, style="Metadata")

    def _add_voice_guide_section(self, doc: "Document", voice_guide_content: str):
        """Add brand voice guide as appendix"""
        doc.add_page_break()

        # Section header
        doc.add_heading("Appendix A: Brand Voice Guide", level=1)

        # Parse markdown content and add to doc
        # For simplicity, add as plain text (can enhance later)
        for line in voice_guide_content.split("\n"):
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("- "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif line.strip().startswith("**") and line.strip().endswith("**"):
                # Bold paragraph
                para = doc.add_paragraph()
                run = para.add_run(line.strip().strip("**"))
                run.bold = True
            elif line.strip():
                doc.add_paragraph(line.strip())

    def _add_schedule_section(self, doc: "Document", schedule_content: str):
        """Add posting schedule as appendix"""
        doc.add_page_break()

        # Section header
        doc.add_heading("Appendix B: Posting Schedule", level=1)

        # Parse markdown content and add to doc
        for line in schedule_content.split("\n"):
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("├─") or line.startswith("└─"):
                # Schedule entry - add as bullet
                doc.add_paragraph(line[3:].strip(), style="List Bullet")
            elif line.strip():
                doc.add_paragraph(line.strip())

    def _add_qa_summary_section(self, doc: "Document", qa_report: QAReport):
        """Add QA summary as appendix"""
        # Lazy import docx types
        from docx.shared import RGBColor

        doc.add_page_break()

        # Section header
        doc.add_heading("Appendix C: Quality Assurance Summary", level=1)

        # Overall status
        status_icon = "✓" if qa_report.overall_passed else "⚠"
        status = doc.add_paragraph()
        status.add_run(f"{status_icon} Overall Status: ")
        status_run = status.add_run("PASSED" if qa_report.overall_passed else "NEEDS ATTENTION")
        status_run.bold = True
        if not qa_report.overall_passed:
            status_run.font.color.rgb = RGBColor(231, 76, 60)  # Red
        else:
            status_run.font.color.rgb = RGBColor(46, 204, 113)  # Green

        # Quality score
        score = doc.add_paragraph()
        score.add_run("Quality Score: ")
        score_run = score.add_run(f"{qa_report.quality_score * 100:.1f}%")
        score_run.bold = True

        doc.add_paragraph()

        # Validation results
        doc.add_heading("Validation Results", level=2)

        # Hook uniqueness
        if qa_report.hook_validation:
            hook_result = qa_report.hook_validation
            doc.add_paragraph(
                f"• Hook Uniqueness: {hook_result.get('uniqueness_score', 0) * 100:.1f}% "
                f"({hook_result.get('metric', '')})",
                style="List Bullet",
            )

        # CTA variety
        if qa_report.cta_validation:
            cta_result = qa_report.cta_validation
            doc.add_paragraph(
                f"• CTA Variety: {cta_result.get('variety_score', 0) * 100:.1f}% "
                f"({cta_result.get('metric', '')})",
                style="List Bullet",
            )

        # Length validation
        if qa_report.length_validation:
            length_result = qa_report.length_validation
            doc.add_paragraph(
                f"• Post Length: {length_result.get('average_length', 0):.0f} words average "
                f"({length_result.get('optimal_ratio', 0) * 100:.0f}% in optimal range)",
                style="List Bullet",
            )

        # Issues
        if qa_report.all_issues:
            doc.add_paragraph()
            doc.add_heading("Issues to Review", level=2)

            for issue in qa_report.all_issues[:5]:  # Show first 5 issues
                doc.add_paragraph(f"• {issue}", style="List Bullet")

            if len(qa_report.all_issues) > 5:
                doc.add_paragraph(
                    f"... and {len(qa_report.all_issues) - 5} more issues",
                    style="Metadata",
                )


# Default generator instance (lazy loaded)
default_docx_generator = None


def get_default_docx_generator() -> DOCXGenerator:
    """Get or create default DOCX generator instance"""
    global default_docx_generator
    if default_docx_generator is None:
        default_docx_generator = DOCXGenerator()
    return default_docx_generator
