"""Briefs router"""

import asyncio
import io
import re
import sys
import time
from pathlib import Path
from typing import Literal

import docx  # python-docx

from fastapi import (
    APIRouter,
    Depends,
    File,
    Header,
    HTTPException,
    Request,
    UploadFile,
    status,
)
from backend.middleware.auth_dependency import get_current_user
from backend.middleware.authorization import (
    verify_brief_ownership,
)  # TR-021: Authorization
from backend.schemas.brief import BriefCreate, BriefResponse
from backend.services import crud
from sqlalchemy.orm import Session

from backend.config import settings
from backend.database import get_db
from backend.models import User
from backend.models.brief_import import FieldExtraction, ParsedBriefResponse
from backend.utils.logger import logger
from backend.utils.http_rate_limiter import standard_limiter, lenient_limiter
from src.validators.prompt_injection_defense import sanitize_prompt_input

router = APIRouter()

# Maximum raw file sizes for brief uploads/parsing
_MAX_TEXT_BYTES = 51_200  # 50 KB for .txt / .md
_MAX_DOCX_BYTES = 5_242_880  # 5 MB for .docx (binary container; text content is small)


def _strip_markdown(text: str) -> str:
    """
    Convert Markdown-formatted text to plain text.

    Preserves all readable content while removing syntax characters that
    interfere with Claude's field extraction (headings, bold, links, etc.).
    """
    # Fenced code blocks — keep content, drop fence markers
    text = re.sub(r"```[^\n]*\n(.*?)```", r"\1", text, flags=re.DOTALL)
    # Inline code — keep content
    text = re.sub(r"`([^`]+)`", r"\1", text)
    # Headings — strip leading # characters
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
    # Bold / italic (**, *, __, _)
    text = re.sub(r"\*{1,3}([^*\n]+)\*{1,3}", r"\1", text)
    text = re.sub(r"_{1,3}([^_\n]+)_{1,3}", r"\1", text)
    # Images ![alt](url) — keep alt text
    text = re.sub(r"!\[([^\]]*)\]\([^)]*\)", r"\1", text)
    # Links [text](url) — keep link text
    text = re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", text)
    # Blockquotes
    text = re.sub(r"^>\s?", "", text, flags=re.MULTILINE)
    # Horizontal rules
    text = re.sub(r"^[-*_]{3,}\s*$", "", text, flags=re.MULTILINE)
    # Unordered list markers (-, *, +)
    text = re.sub(r"^[ \t]*[-*+]\s+", "", text, flags=re.MULTILINE)
    # Ordered list markers (1. 2. etc.)
    text = re.sub(r"^[ \t]*\d+\.\s+", "", text, flags=re.MULTILINE)
    # HTML tags
    text = re.sub(r"<[^>]+>", "", text)
    # Collapse 3+ blank lines to 2
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _extract_brief_text(content: bytes, file_ext: str) -> str:
    """
    Extract plain text from a brief file.

    For .txt files the bytes are decoded as UTF-8 and returned as-is.
    For .md files the bytes are decoded then Markdown syntax is stripped
    so Claude receives clean prose rather than raw formatting characters.
    For .docx files python-docx is used to extract paragraph and table text.

    Args:
        content:  Raw file bytes.
        file_ext: Lowercase file extension including the dot (e.g. ".docx").

    Returns:
        Extracted plain-text string.

    Raises:
        ValueError: If the file cannot be decoded or parsed.
    """
    if file_ext == ".docx":
        try:
            document = docx.Document(io.BytesIO(content))
        except Exception as exc:
            raise ValueError(f"Could not read .docx file: {exc}") from exc

        lines: list[str] = []
        for para in document.paragraphs:
            text = para.text.strip()
            if text:
                lines.append(text)
        for table in document.tables:
            for row in table.rows:
                row_text = "\t".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    lines.append(row_text)
        return "\n".join(lines)

    # Decode UTF-8 for .txt and .md
    try:
        raw_text = content.decode("utf-8")
    except UnicodeDecodeError as exc:
        raise ValueError(f"File must be UTF-8 encoded: {exc}") from exc

    if file_ext == ".md":
        stripped = _strip_markdown(raw_text)
        logger.info(f"Stripped Markdown: {len(raw_text)} chars → {len(stripped)} chars plain text")
        return stripped

    return raw_text


@router.post("/create", response_model=BriefResponse, status_code=status.HTTP_201_CREATED)
@standard_limiter.limit("100/hour")  # TR-004: Standard operation
async def create_brief_from_text(
    request: Request,
    brief: BriefCreate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Create brief from pasted text.

    Rate limit: 100/hour per IP+user (standard operation)
    Authorization: TR-021 - User must own the project
    """
    # Verify project exists
    project = crud.get_project(db, brief.project_id)
    if not project:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Project not found")

    # Reattach detached object to session for attribute access
    project = db.merge(project)

    # TR-021: Verify user owns the project before creating brief
    if project.user_id != current_user.id and not current_user.is_superuser:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Access denied: You don't own this project",
        )

    # Check if brief already exists for this project
    existing_brief = crud.get_brief_by_project(db, brief.project_id)
    if existing_brief:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Brief already exists for this project",
        )

    # SECURITY (TR-020): Sanitize brief content before saving (will be passed to LLM later)
    try:
        sanitized_content = sanitize_prompt_input(brief.content, strict=False)
        logger.info(f"Sanitized brief content for project {brief.project_id}")
    except ValueError as e:
        logger.warning(f"Prompt injection detected in brief content: {e}")
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Brief content contains potentially unsafe patterns. Please review and resubmit.",
        )

    # Save brief to file (use sanitized content)
    briefs_dir = Path(settings.BRIEFS_DIR)
    briefs_dir.mkdir(parents=True, exist_ok=True)
    file_path = briefs_dir / f"{brief.project_id}.txt"
    file_path.write_text(sanitized_content, encoding="utf-8")

    # Create brief with sanitized content
    sanitized_brief = BriefCreate(project_id=brief.project_id, content=sanitized_content)
    return crud.create_brief(db, sanitized_brief, source="paste", file_path=str(file_path))


@router.post("/upload", response_model=BriefResponse, status_code=status.HTTP_201_CREATED)
@standard_limiter.limit("100/hour")  # TR-004: Standard operation
async def upload_brief_file(
    request: Request,
    project_id: str,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Upload brief file.

    Rate limit: 100/hour per IP+user (standard operation)
    Authorization: TR-021 - User must own the project
    """
    # Verify project exists
    project = crud.get_project(db, project_id)
    if not project:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Project not found")

    # Reattach detached object to session for attribute access
    project = db.merge(project)

    # TR-021: Verify user owns the project before uploading brief
    if project.user_id != current_user.id and not current_user.is_superuser:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Access denied: You don't own this project",
        )

    # Check file extension
    file_ext = Path(file.filename or "").suffix.lower()
    if file_ext not in settings.allowed_extensions_list:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Invalid file type. Allowed: {settings.ALLOWED_BRIEF_EXTENSIONS}",
        )

    # Read and extract file content
    content = await file.read()
    try:
        text_content = _extract_brief_text(content, file_ext)
    except ValueError as e:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=str(e),
        )

    # SECURITY (TR-020): Sanitize uploaded brief content before saving
    try:
        sanitized_content = sanitize_prompt_input(text_content, strict=False)
        logger.info(f"Sanitized uploaded brief for project {project_id}")
    except ValueError as e:
        logger.warning(f"Prompt injection detected in uploaded brief: {e}")
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Uploaded file contains potentially unsafe patterns. Please review and resubmit.",
        )

    # Save file (use sanitized content)
    briefs_dir = Path(settings.BRIEFS_DIR)
    briefs_dir.mkdir(parents=True, exist_ok=True)
    file_path = briefs_dir / f"{project_id}{file_ext}"
    file_path.write_text(sanitized_content, encoding="utf-8")

    # Create brief record (use sanitized content)
    brief_data = BriefCreate(project_id=project_id, content=sanitized_content)
    return crud.create_brief(db, brief_data, source="upload", file_path=str(file_path))


@router.get("/{brief_id}", response_model=BriefResponse)
@lenient_limiter.limit("1000/hour")  # TR-004: Cheap read operation
async def get_brief(
    request: Request,
    brief_id: str,
    brief=Depends(verify_brief_ownership),  # TR-021: Authorization check
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Get brief by ID.

    Rate limit: 1000/hour (cheap read operation)
    Authorization: TR-021 - User must own brief's project
    """
    # TR-021: brief already verified by dependency
    return brief


@router.post("/parse", response_model=ParsedBriefResponse)
@standard_limiter.limit("100/hour")  # TR-004: AI parsing operation (moderate cost)
async def parse_brief_file(
    request: Request,
    file: UploadFile = File(...),
    authorization: str | None = Header(default=None),
):
    """
    Parse uploaded brief file and extract client data with confidence scores.

    Rate limit: 100/hour per IP+user (AI parsing operation)
    """
    start_time = time.time()

    if not authorization:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Not authenticated",
            headers={"WWW-Authenticate": "Bearer"},
        )

    # Validate file extension (.txt, .md, .docx)
    file_ext = Path(file.filename or "").suffix.lower()
    if file_ext not in [".txt", ".md", ".docx"]:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail={
                "code": "INVALID_FILE_TYPE",
                "message": "Only .txt and .md files are supported",
                "details": {"filename": file.filename, "extension": file_ext},
            },
        )

    # Validate file size — .docx is a binary container so its raw size is larger
    content = await file.read()
    file_size = len(content)
    max_size = _MAX_DOCX_BYTES if file_ext == ".docx" else _MAX_TEXT_BYTES
    max_label = "5MB" if file_ext == ".docx" else "50KB"

    if file_size > max_size:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail={
                "code": "FILE_TOO_LARGE",
                "message": f"File must be less than {max_label}",
                "details": {
                    "filename": file.filename,
                    "sizeBytes": file_size,
                    "maxSizeBytes": max_size,
                },
            },
        )

    # Extract plain text from the file
    try:
        text_content = _extract_brief_text(content, file_ext)
    except ValueError as e:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail={
                "code": "ENCODING_ERROR",
                "message": str(e),
                "details": {"filename": file.filename},
            },
        )

    # SECURITY (TR-020): Sanitize brief content before parsing with LLM
    try:
        sanitized_content = sanitize_prompt_input(text_content, strict=False)
        logger.info(f"Sanitized brief content for parsing: {file.filename}")
    except ValueError as e:
        logger.warning(f"Prompt injection detected in brief file: {e}")
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail={
                "code": "SECURITY_VALIDATION_FAILED",
                "message": "File contains potentially unsafe content patterns",
                "details": {"filename": file.filename, "error": str(e)},
            },
        )

    # Parse with BriefParserAgent: run 3 attempts in parallel, keep the best result.
    # parse_brief() is synchronous; run_in_executor moves each call to a thread so all
    # three Claude API requests are in-flight concurrently.
    try:
        # Import here to avoid circular dependency
        sys.path.insert(0, str(Path(__file__).parent.parent.parent))
        from src.agents.brief_parser import BriefParserAgent

        parser = BriefParserAgent()
        loop = asyncio.get_event_loop()
        attempts = await asyncio.gather(
            *[loop.run_in_executor(None, parser.parse_brief, sanitized_content) for _ in range(3)],
            return_exceptions=True,
        )

        # Score each successful attempt; discard failures
        best_fields = None
        best_score = -1
        for attempt in attempts:
            if isinstance(attempt, Exception):
                logger.warning(f"Parse attempt failed (will use best of remaining): {attempt}")
                continue
            fields = _add_confidence_scores(attempt, text_content)
            score = _score_fields(fields)
            logger.info(f"Parse attempt score: {score}")
            if score > best_score:
                best_score = score
                best_fields = fields

        if best_fields is None:
            raise RuntimeError("All 3 parse attempts failed")

        # Generate warnings for missing/low-confidence fields
        warnings = _generate_warnings(best_fields)

        # Calculate metadata
        parse_time_ms = int((time.time() - start_time) * 1000)
        fields_extracted = sum(1 for field in best_fields.values() if field.confidence != "low")
        fields_total = len(best_fields)

        return ParsedBriefResponse(
            success=True,
            fields=best_fields,
            warnings=warnings,
            metadata={
                "filename": file.filename,
                "parseTimeMs": parse_time_ms,
                "fieldsExtracted": fields_extracted,
                "fieldsTotal": fields_total,
            },
        )

    except Exception as e:
        logger.error(f"Brief parsing failed: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail={
                "code": "PARSE_FAILED",
                "message": "Failed to parse brief file",
                "details": {"filename": file.filename, "error": str(e)},
            },
        )


def _score_fields(fields: dict) -> int:
    """Score extracted fields by confidence quality to select the best parse attempt.

    Weights: high = 3, medium = 1, low = 0.
    A result with more high-confidence fields beats one with many medium fields.
    """
    weights = {"high": 3, "medium": 1, "low": 0}
    return sum(weights.get(f.confidence, 0) for f in fields.values())


def _add_confidence_scores(parsed_brief, original_text: str) -> dict:
    """Add confidence scores to extracted fields based on content quality"""
    fields = {}

    # Map ClientBrief fields to extraction results
    field_mapping = {
        "companyName": parsed_brief.company_name,
        "founderName": parsed_brief.founder_name,
        "businessDescription": parsed_brief.business_description,
        "industry": parsed_brief.industry,
        "keywords": parsed_brief.keywords,
        "competitors": parsed_brief.competitors,
        "location": parsed_brief.location,
        "idealCustomer": parsed_brief.ideal_customer,
        "mainProblemSolved": parsed_brief.main_problem_solved,
        "tonePreference": (
            parsed_brief.tone_preference.value if parsed_brief.tone_preference else "professional"
        ),
        "toneToAvoid": parsed_brief.tone_to_avoid,
        "brandPersonality": parsed_brief.brand_personality,
        "platforms": (
            [p.value for p in parsed_brief.target_platforms]
            if parsed_brief.target_platforms
            else []
        ),
        "postingFrequency": parsed_brief.posting_frequency,
        "dataUsage": (parsed_brief.data_usage.value if parsed_brief.data_usage else "moderate"),
        "customerPainPoints": parsed_brief.customer_pain_points,
        "customerQuestions": parsed_brief.customer_questions,
        "mainCta": parsed_brief.main_cta,
        "measurableResults": parsed_brief.measurable_results,
        "stories": parsed_brief.stories,
        "misconceptions": parsed_brief.misconceptions,
        "keyPhrases": parsed_brief.key_phrases,
    }

    for field_name, field_value in field_mapping.items():
        # Determine confidence based on field completeness
        confidence: Literal["high", "medium", "low"]
        value: str | list | None
        if field_value is None or field_value == "" or field_value == []:
            confidence = "low"
            value = None
        elif isinstance(field_value, str):
            # String fields: high if >10 chars, medium if 3-10, low otherwise
            if len(field_value) > 10:
                confidence = "high"
            elif len(field_value) >= 3:
                confidence = "medium"
            else:
                confidence = "low"
            value = field_value if field_value else None
        elif isinstance(field_value, list):
            # List fields: keep as list so the frontend can spread them correctly.
            # high if 2+ items, medium if 1 item, low if empty.
            if len(field_value) >= 2:
                confidence = "high"
            elif len(field_value) == 1:
                confidence = "medium"
            else:
                confidence = "low"
            value = field_value if field_value else None
        else:
            # Other types (enums, etc.)
            confidence = "high" if field_value else "low"
            value = field_value

        # Try to find source line number (approximate) — only for string values
        source = None
        if value and isinstance(value, str):
            lines = original_text.split("\n")
            for i, line in enumerate(lines, 1):
                if value[:20] in line:  # Match first 20 chars
                    source = f"line {i}"
                    break

        fields[field_name] = FieldExtraction(value=value, confidence=confidence, source=source)

    return fields


def _generate_warnings(fields: dict) -> list:
    """Generate warnings for missing or low-confidence fields"""
    warnings = []

    # Required fields that should have high confidence
    required_fields = ["companyName", "businessDescription", "idealCustomer"]

    for field_name in required_fields:
        field = fields.get(field_name)
        if not field or field.confidence == "low":
            warnings.append(f"{field_name} not found or low confidence - may need manual entry")

    # Optional fields with defaults
    tone_field = fields.get("tonePreference")
    if tone_field and tone_field.confidence == "low":
        warnings.append("tonePreference not found, defaulting to 'professional'")

    platforms_field = fields.get("platforms")
    if platforms_field and platforms_field.confidence == "low":
        warnings.append("No platforms specified - will need to select manually")

    # Bug #110: warn about unsupported platforms extracted from the brief
    if platforms_field and platforms_field.value:
        from src.models.client_brief import Platform as _Platform

        supported = {p.value for p in _Platform}
        extracted = platforms_field.value if isinstance(platforms_field.value, list) else []
        unsupported = [p for p in extracted if str(p).lower() not in supported]
        if unsupported:
            names = ", ".join(unsupported)
            warnings.append(
                f"Unsupported platforms detected: {names}. "
                f"Supported platforms are: {', '.join(sorted(supported))}. "
                "These platforms will be ignored during content generation."
            )

    return warnings
